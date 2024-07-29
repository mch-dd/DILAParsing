import os
import tarfile
import json
import xml.etree.ElementTree as ET
from docx import Document
from odf import text, teletype
import shutil
from odf.opendocument import load
from docx.opc.exceptions import PackageNotFoundError
from lxml.etree import XMLSyntaxError  # Import XMLSyntaxError

def extract_text_from_docx(docx_path):
    """Extracts text from a DOCX file with error handling for missing packages and other issues."""
    try:
        doc = Document(docx_path)
        text_content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        word_count = len(text_content.split())
        return text_content, word_count
    except (KeyError, PackageNotFoundError, XMLSyntaxError) as e:
        print(f"Error reading {docx_path}: {e}")
        return "ERROR in reading text", 0  # Return error message and 0 word count on specific errors

def extract_tar_gz(tar_path, extract_path):
    """Extracts a .tar.gz file to a specified directory."""
    try:
        with tarfile.open(tar_path, 'r:gz') as tar:
            tar.extractall(path=extract_path)
        print(f'Extracted {tar_path}')
    except Exception as e:
        print(f'Error extracting {tar_path}: {e}')

def find_files(directory, extension):
    """Finds all files in directory and its subdirectories with the given extension."""
    for root, dirs, files in os.walk(directory):
        for file in files:
            if file.endswith(extension):
                yield os.path.join(root, file)

def convert_odt_to_docx(odt_path, docx_path):
    """Converts an ODT file to DOCX format."""
    textdoc = load(odt_path)
    alltext = ""
    for paragraph in textdoc.getElementsByType(text.P):
        alltext += teletype.extractText(paragraph) + "\n"
    
    doc = Document()
    doc.add_paragraph(alltext)
    doc.save(docx_path)
    print("\033[94m" + f"Converted {odt_path} to DOCX" + "\033[0m")

root_directory = 'ACCO'
limit_to_first_folder = False  # Set this to False to process all folders

for index, folder in enumerate(os.listdir(root_directory)):
    if limit_to_first_folder and index > 0:
        break

    if folder.endswith('.tar.gz'):
        folder_path = os.path.join(root_directory, folder[:-7])  # Remove '.tar.gz'
        extract_tar_gz(os.path.join(root_directory, folder), folder_path)

        xml_files = list(find_files(folder_path, '.xml'))  # Convert generator to list to count files
        total_xml_files = len(xml_files)
        print(f"Found {total_xml_files} XML files in {folder_path}.")
        
        json_data = []
        for processed_count, xml_file in enumerate(xml_files, start=1):
            tree = ET.parse(xml_file)
            root = tree.getroot()
            # The rest of the XML processing code remains unchanged...
            
            # After processing an XML file, update progress
            progress = (processed_count / total_xml_files) * 100
            print(f"Processed {processed_count}/{total_xml_files} XML files. Progress: {progress:.2f}%")

            if processed_count >= total_xml_files:
                print("\033[92m" + f"All XML files processed. Saving data to JSON..." + "\033[0m")

            # Extract relevant information
            entry = {
                "ID": root.find(".//ID").text if root.find(".//ID") is not None else "",
                "Date_effet": root.find(".//DATE_EFFET").text if root.find(".//DATE_EFFET") is not None else "",
                "Date_fin": root.find(".//DATE_FIN").text if root.find(".//DATE_FIN") is not None else "",
                "SIRET": root.find(".//SIRET").text if root.find(".//SIRET") is not None else "",
                "Raison_sociale": root.find(".//RAISON_SOCIALE").text if root.find(".//RAISON_SOCIALE") is not None else "",
                "Themes": [{"Code": theme.find("CODE").text, "Libelle": theme.find("LIBELLE").text, "Groupe": theme.find("GROUPE").text} for theme in root.findall(".//THEME")],
                "Text": ""
            }
            document_name = root.find(".//DOCUMENT_BUREAUTIQUE").text.split('/')[-1] if root.find(".//DOCUMENT_BUREAUTIQUE") is not None else ""
            for file_path in find_files(folder_path, document_name):
                if file_path.endswith('.docx') or file_path.endswith('.odt'):
                    if file_path.endswith('.odt'):
                        temp_docx_path = file_path + '.docx'
                        convert_odt_to_docx(file_path, temp_docx_path)
                        file_path = temp_docx_path  # Update file_path to point to the converted DOCX
                        
                    text_content, word_count = extract_text_from_docx(file_path)
                    entry["Text"] = text_content
                    entry["Word_count"] = word_count

                    if file_path.endswith('.odt'):  # Clean up temporary DOCX file if it was an ODT conversion
                        os.remove(temp_docx_path)

            json_data.append(entry)
        # Save to JSON file
        json_path = os.path.join(folder_path, folder[:-7] + '.json')
        with open(json_path, 'w', encoding='utf-8') as f:
            json.dump(json_data, f, ensure_ascii=False, indent=4)
        print("\033[92m" + f"Successfully saved data to {json_path}" + "\033[0m")
